# -*- coding: utf-8 -*-
"""
面向消费者与潜在合作分销伙伴的说明文档（.docx）。
文风：商业正式、通俗易懂，不涉及后台技术实现细节。
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt


def set_run_font(run, size_pt=11):
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size_pt)


def set_heading_run_font(run, size_pt=14):
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size_pt)


def add_para(doc, text, size_pt=11):
    p = doc.add_paragraph(text)
    for r in p.runs:
        set_run_font(r, size_pt)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        set_heading_run_font(r, 16 if level == 1 else 13)
    return h


def _style_table_cell(cell, text, size_pt=10):
    cell.text = text
    for p in cell.paragraphs:
        for r in p.runs:
            set_run_font(r, size_pt)


def add_distributor_level_simple_table(doc):
    """
    分销等级对照表：至少展示两档差异（示例两行 + 预留行），列上突出「相对基础档」的差异维度。
    """
    headers = [
        "等级名称",
        "档位排序",
        "推广与分润\n（相对基础档）",
        "供货成本与团队激励\n（相对基础档）",
        "备注",
    ]
    # 1 表头 + 2 行「两档差异」示例 + 2 行预留
    example_rows = [
        [
            "基础分销档（示例）",
            "入门档",
            "以直接分享、促成下单为主；享受基础推广奖励与常规分享比例，适合起步伙伴。",
            "对应常规供货或成本政策；团队激励按入门规则参与，侧重学习与跑通流程。",
            "示例仅供理解结构；名称、比例以平台公示为准。",
        ],
        [
            "进阶分销档（示例）",
            "高于入门档",
            "在入门档之上，通常享有更高分享比例或更大单笔可分润空间（是否叠加活动以公示为准）。",
            "通常对应更优惠的供货或成本条件，并可能获得更高权重的团队激励与渠道支持。",
            "示例仅供理解结构；名称、比例以平台公示为准。",
        ],
        ["（平台自填等级）", "—", "—", "—", "可写升级条件、有效期等。"],
        ["（平台自填等级）", "—", "—", "—", "—"],
    ]
    table = doc.add_table(rows=1 + len(example_rows), cols=len(headers))
    table.style = "Table Grid"
    for c, h in enumerate(headers):
        _style_table_cell(table.rows[0].cells[c], h, 9)
    for r, row_cells in enumerate(example_rows, start=1):
        for c, val in enumerate(row_cells):
            _style_table_cell(table.rows[r].cells[c], val, 9)

    add_heading(doc, "表后说明：各列含义", level=3)
    col_help = [
        "「等级名称」：您在小程序或合作资料中看到的分销身份名称；上表前两行为**对照示例**，展示「入门档」与「更高一档」在表述上的典型差异，便于读者理解；**非对贵司实际等级的承诺**，贵司应将示例替换为真实等级名称。",
        "「档位排序」：标明各档之间相对高低或先后关系（如入门、进阶等），具体以贵司规则为准。",
        "「推广与分润（相对基础档）」：用一句话概括该档在**直接/间接推广奖励、分享比例**等方面，相对最低档或基础档**多了什么、优在哪里**；不写具体公式。",
        "「供货成本与团队激励（相对基础档）」：概括该档在**进货/成本政策**以及**团队激励权重或资格**上，相对基础档**有何不同**；具体数值以合同、后台或活动页为准。",
        "「备注」：可写升级条件、考核周期、区域限制、与活动叠加以何者优先等；无则填横线。",
    ]
    for line in col_help:
        p = doc.add_paragraph(line, style="List Bullet")
        for r in p.runs:
            set_run_font(r, 10)


def build_document(path: Path) -> None:
    doc = Document()
    add_heading(doc, "分销奖励与收益说明", level=1)

    add_para(
        doc,
        "尊敬的用户与合作伙伴：感谢您关注本平台的分享与分销计划。本文旨在以清晰、务实的方式，说明参与推广后您可能获得的收益类型（含推广佣金与团队激励等）、大致来源，以及收益如何进入您的账户并用于消费或提现。具体比例、活动规则以平台届时公示或协议约定为准。",
    )

    add_heading(doc, "一、计划定位", level=2)
    add_para(
        doc,
        "本计划鼓励用户通过正当分享，将优质商品或服务推荐给亲友与社交圈。在符合平台规则的前提下，订单成交后，平台将按既定规则向相关参与方分配一定奖励，以体现对推广与渠道贡献的认可。",
    )

    add_heading(doc, "二、哪些人可能获得奖励", level=2)
    roles = [
        "直接分享人：通过您的分享链接、海报或推荐关系完成注册或下单的，您作为直接推荐人，有机会获得与「直接推广」相关的奖励。",
        "间接推荐链条中的参与者：在平台允许的多级推荐关系内，符合条件的上级伙伴，可能依据规则获得与团队推广、等级差异等相关的补充性奖励（若当前业务已启用相应功能）。",
        "具备分销身份的合作伙伴：若您已申请并通过平台审核、取得相应合作等级，还可依据等级所对应的经营政策，在订单完成后参与利润分配。具体身份与权益以您账户内展示及合同约定为准。",
    ]
    for t in roles:
        p = doc.add_paragraph(t, style="List Bullet")
        for r in p.runs:
            set_run_font(r)

    add_heading(doc, "三、分销等级一览（简表）", level=2)
    add_para(
        doc,
        "为便于理解**至少两档分销身份**在权益上的典型差异，下表采用「对照式」列设计：除等级名称与排序外，专门从「推广与分润」「供货成本与团队激励」两个维度，说明**相对基础档**多出的价值。表中**前两行为结构示例**，贵司应替换为真实等级名称与表述；后两行为预留行，可增列更多档位。若与线上一致，读者也可直接对照小程序内展示。",
    )
    add_distributor_level_simple_table(doc)

    add_heading(doc, "四、奖励从哪里来", level=2)
    add_para(
        doc,
        "奖励来源于订单在扣除商品成本、平台基础运营成本及依法应由商户承担的费用之后，预留用于市场拓展与渠道激励的部分。并非每一笔订单都会产生全部类型的奖励，是否发放、发放对象与金额，取决于订单是否有效完成、是否属于参与计奖的商品范围、以及当时适用的活动与等级政策。",
    )

    add_heading(doc, "五、分配方式（原则说明）", level=2)
    add_para(
        doc,
        "为兼顾「鼓励分享」与「渠道公平」，平台通常按以下思路分配（实际以系统执行与公示规则为准）：优先保障直接分享所产生的推广激励；在存在多级合作结构时，再按等级与团队贡献规则，在剩余空间内进行二次分配。您无需自行计算，订单完成后可在「佣金/奖励」相关页面查看明细。",
    )

    add_heading(doc, "六、团队激励说明", level=2)
    add_para(
        doc,
        "在推广佣金之外，若您已具备平台认定的团队拓展身份或等级，且相关功能已开启，您还有机会获得「团队激励」类奖励，用于回报您在团队培育、动销带动等方面的贡献。该类奖励与单笔订单的推广佣金性质不同，平台可能在账户中单独展示为「团队激励余额」或类似名称，以便您区分管理。",
    )
    add_para(
        doc,
        "团队激励的常见形式包括：其一，与单笔订单关联、在订单完成并经平台确认后计入的激励（您可在佣金/奖励明细中查看来源订单）；其二，按自然月汇总业绩后统一核算的激励（若平台已启用月度结算，请关注每月公示或到账提醒）。是否同时适用两种形式、具体比例与参与条件，以您账户内展示及当期活动规则为准。",
    )
    add_para(
        doc,
        "关于使用方式：团队激励余额与「可提现佣金」在系统中通常分列管理。若平台允许将团队激励用于商城抵扣、提现或转入其他权益，将以收银台、提现页及公告为准；若某渠道暂未开放，亦请以页面提示为准，避免误解。",
    )

    add_heading(doc, "七、收益如何落实到您手中", level=2)
    add_para(
        doc,
        "经平台确认后的奖励，会计入您账户中的「可提现佣金」或同类余额字段，您可通过以下方式使用：",
    )
    use_bullets = [
        "在本平台商城下单支付时，可使用账户内的佣金余额抵扣应付金额。佣金抵扣消费的部分，平台不另行收取手续费，与现金支付享受同等商品与服务（具体是否支持抵扣、单笔上限以收银台展示为准）。",
        "申请提现至微信钱包或银行卡：您发起提现的金额，为从可提现佣金中扣减的申请总额。根据平台规则，可能对每笔提现收取一定手续费（例如按提现金额的一定比例和/或每笔固定金额收取），手续费从申请金额内扣除；扣除后剩余部分为实际划付给您的金额。微信或银行等第三方渠道若另有规定，以该渠道说明为准。",
        "提现需满足平台设定的最低金额等条件；提交后进入审核与打款流程，请留意站内消息或短信通知。",
    ]
    for t in use_bullets:
        p = doc.add_paragraph(t, style="List Bullet")
        for r in p.runs:
            set_run_font(r)

    add_heading(doc, "八、关于提现手续费与到账金额", level=2)
    add_para(
        doc,
        "在您申请将推广佣金提现至微信、银行等外部账户时，平台需依法配合涉税信息管理、代扣代缴或申报辅导，并承担支付通道、对账、反洗钱合规及账务处理等综合成本。**因此，平台所设提现手续费，首要用途为覆盖与提现相关的税费及上述合规与运营支出**；具体费率由平台在后台配置（常见为按提现金额的一定比例和/或每笔固定金额），二者可同时存在。",
    )
    add_para(
        doc,
        "举例：若您申请提现 100 元，手续费合计 2 元，则实际划付至您外部账户的金额为 98 元；该 100 元仍从您的可提现佣金余额中扣减。若手续费规则导致到账金额低于法定或平台最低打款标准，系统将提示您调整申请金额。手续费不构成对商品或服务的加价，亦不作为平台对您的经营性加价利润。",
    )
    add_para(
        doc,
        "再次说明：上述手续费仅针对「提现至外部账户」的行为；使用佣金在本平台直接抵扣购物款时，不收取该提现类手续费。",
    )

    add_heading(doc, "九、合规与诚信提示", level=2)
    add_para(
        doc,
        "请遵守国家广告法、反不正当竞争及平台关于真实宣传、禁止传销等规定。奖励属于个人经营或劳务所得的，请依法履行纳税申报义务。平台有权根据监管要求与经营需要调整规则，并将通过合理方式提前或及时公示。",
    )

    add_heading(doc, "十、进一步了解", level=2)
    add_para(
        doc,
        "若您需要了解当前等级、可提现余额、单笔抵扣上限或提现手续费标准，请登录小程序或联系平台客服，以实时展示与人工答复为准。",
    )

    doc.core_properties.title = "分销奖励与收益说明"
    doc.core_properties.subject = "消费者与合作伙伴"
    doc.save(path)


def main():
    out = Path("e:/MMBS16/docs/分销奖励与收益说明.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    try:
        build_document(out)
        print("written:", out)
    except PermissionError:
        alt = out.with_name("分销奖励与收益说明_生成版.docx")
        build_document(alt)
        print("written (原文件被占用，已写入):", alt)
        print("请关闭 Word 后删除或替换原文件，再将本文件重命名为：分销奖励与收益说明.docx")


if __name__ == "__main__":
    main()
